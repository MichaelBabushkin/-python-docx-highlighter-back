from time import sleep
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
import os
import re



def highlight_translations(file,fname):
    document = Document(file)
    print(document)
    for para in document.paragraphs:
        for run in para.runs:
            if ':' in run.text:
                row = run.text.split(':')
                run.clear()
                print ("row" , row)
                for i in range(len(row)-1):
                    run.add_text(row[0])
                    run.add_text(':')
                    if '{{' not in row[1]:
                        para.add_run(row[1]).font.highlight_color = WD_COLOR_INDEX.YELLOW
                    else:
                        text_in_brackets = re.findall('{(.+?)}',row[1])
                        words = row[1].split(' ')
                        for word in words:
                            print ("word" , word)
                            if word.startswith('{{'):
                                para.add_run(word + " ").font.highlight_color = WD_COLOR_INDEX.WHITE    
                            else:
                                para.add_run(word + " ").font.highlight_color = WD_COLOR_INDEX.YELLOW
    if not os.path.exists(os.path.normpath(os.path.expanduser("~/highlighted_files"))):
        os.makedirs(os.path.normpath(os.path.expanduser("~/highlighted_files")))
    save_path = os.path.normpath(os.path.expanduser("~/highlighted_files"))
    filepath = os.path.join(save_path, fname+'_translated.docx')
    document.save(filepath)
    sleep(2)
    return {'status':200, "filepath":save_path}